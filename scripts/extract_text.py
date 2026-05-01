# SPDX-License-Identifier: GPL-3.0-or-later
# Copyright (C) 2026 Patrick R. Wallace, Hamilton College LITS
#
# This program is free software: you can redistribute it and/or modify it
# under the terms of the GNU General Public License as published by the Free
# Software Foundation, either version 3 of the License, or (at your option)
# any later version.  See <https://www.gnu.org/licenses/gpl-3.0.html>.
#
# AI ASSISTANCE NOTICE: Developed with assistance from Claude (Anthropic).
# Reviewed and tested; verify behavior in your own environment.

"""
extract_text.py — Extract plain text from native (non-scanned) document types.

Supported formats and their backends
--------------------------------------
  .pdf          pdfminer.six   Scanned PDFs (no text layer) are detected and
                               flagged for re-routing to ocr.py.
  .docx, .odt   python-docx    Paragraph text; core properties as metadata.
  .html, .htm   BeautifulSoup  Script/style stripped; visible text only.
  .csv          stdlib csv     Rendered as pipe-separated plain text.
  .txt, .md     built-in       Direct UTF-8 read.

Return value
------------
All functions return a dict with these keys:

    text      str        Extracted plain text (stripped of leading/trailing
                         whitespace).
    type      str        File type derived from extension (e.g. "pdf").
    source    Path       Original file path.
    metadata  dict       Any available metadata (title, author, etc.).
    error     str|None   Error message if extraction failed; None on success.
    is_scan   bool       True if the file appears to be a scanned PDF with no
                         usable text layer.  Only ever True for .pdf files.

The function never raises — errors are captured in result["error"] so callers
can decide how to handle failures (log and continue, quarantine, etc.).

CLI exit codes
--------------
    0  Success — text printed to stdout.
    1  Extraction error — message on stderr.
    2  Scanned PDF detected — message on stderr; caller should use ocr.py.

Usage (CLI)
-----------
    python scripts/extract_text.py path/to/document.pdf

Usage (import)
--------------
    from scripts.extract_text import extract_text
    result = extract_text(Path("paper.pdf"))
    if result["is_scan"]:
        # re-route to ocr.py
    elif result["error"]:
        # handle failure
    else:
        text = result["text"]
"""

import csv
import sys
from io import StringIO
from pathlib import Path


# ---------------------------------------------------------------------------
# Per-format extractors
# ---------------------------------------------------------------------------

def _extract_pdf(filepath: Path) -> tuple[str, dict, bool]:
    """Extract text from a PDF that has a native text layer.

    Returns (text, metadata, is_scan).

    is_scan is set to True when the extracted text is suspiciously short —
    fewer than 100 non-whitespace characters across the entire document.
    This is a reliable heuristic for scanned PDFs where the PDF container
    holds page images but no embedded text.  The threshold is deliberately
    conservative: a real single-page document will almost always exceed 100
    non-whitespace characters.
    """
    from pdfminer.high_level import extract_text as pdfminer_extract
    from pdfminer.pdfdocument import PDFDocument
    from pdfminer.pdfparser import PDFParser

    text = pdfminer_extract(str(filepath))

    # Attempt to read standard PDF metadata from the document info dictionary.
    # This is best-effort: malformed or encrypted PDFs may not expose metadata,
    # and we must not let a metadata failure abort the text extraction.
    metadata = {}
    try:
        with open(filepath, "rb") as f:
            parser = PDFParser(f)
            doc = PDFDocument(parser)
            if doc.info:
                raw = doc.info[0]
                for key in ("Title", "Author", "Subject", "Keywords"):
                    value = raw.get(key)
                    if value:
                        # PDF metadata values are bytes in older documents.
                        if isinstance(value, bytes):
                            value = value.decode("utf-8", errors="replace")
                        metadata[key.lower()] = value.strip()
    except Exception:
        pass

    # Count non-whitespace characters to detect scanned (image-only) PDFs.
    non_ws_count = len(text.replace(" ", "").replace("\n", "").replace("\t", ""))
    is_scan = non_ws_count < 100

    return text, metadata, is_scan


def _extract_docx(filepath: Path) -> tuple[str, dict]:
    """Extract text from a .docx or .odt file using python-docx.

    Paragraphs are joined with blank lines between them to preserve the
    visual structure of the document.  Empty paragraphs (section breaks,
    spacing) are filtered out.
    """
    from docx import Document
    doc = Document(str(filepath))

    # Filter out truly empty paragraphs; preserve intentional blank lines
    # between content paragraphs by using double newline as the separator.
    text = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())

    metadata = {}
    props = doc.core_properties
    if props.title:
        metadata["title"] = props.title
    if props.author:
        metadata["author"] = props.author

    return text, metadata


def _extract_html(filepath: Path) -> tuple[str, dict]:
    """Extract visible text from an HTML file using BeautifulSoup.

    Removes script, style, and head elements — their text content is not
    meaningful for research purposes.  Uses the lxml parser for speed and
    robust handling of malformed HTML.
    """
    from bs4 import BeautifulSoup

    raw = filepath.read_text(encoding="utf-8", errors="replace")
    soup = BeautifulSoup(raw, "lxml")

    # Remove non-content elements before extracting text.
    for tag in soup(["script", "style", "head"]):
        tag.decompose()

    # get_text with a newline separator, then strip blank lines.
    lines = [line.strip() for line in soup.get_text(separator="\n").splitlines()]
    text = "\n".join(line for line in lines if line)

    # The <title> element is the most reliable source of a document title.
    metadata = {}
    title_tag = soup.find("title")
    if title_tag:
        metadata["title"] = title_tag.get_text(strip=True)

    return text, metadata


def _extract_csv(filepath: Path) -> tuple[str, dict]:
    """Render a CSV file as pipe-separated plain text.

    Pipe separation is chosen over the original delimiter because it is
    human-readable, survives whitespace normalisation, and is easy for Claude
    to parse into a markdown table.  The first row is treated as a header.
    """
    raw = filepath.read_text(encoding="utf-8", errors="replace")
    reader = csv.reader(StringIO(raw))
    rows = list(reader)

    lines = [" | ".join(cell.strip() for cell in row) for row in rows if any(row)]

    metadata = {}
    if rows:
        metadata["columns"] = rows[0]
        metadata["row_count"] = str(len(rows) - 1)

    return "\n".join(lines), metadata


def _extract_plaintext(filepath: Path) -> tuple[str, dict]:
    """Read a plain text or markdown file directly.

    Uses UTF-8 decoding with error replacement so malformed bytes do not
    abort extraction.
    """
    return filepath.read_text(encoding="utf-8", errors="replace"), {}


# ---------------------------------------------------------------------------
# Dispatcher
# ---------------------------------------------------------------------------

# Maps each supported extension to the function that handles it.
# .odt reuses the .docx extractor because python-docx handles both formats
# via the same OOXML-based API.
_EXTRACTORS = {
    ".pdf":  _extract_pdf,      # returns 3-tuple (text, metadata, is_scan)
    ".docx": _extract_docx,
    ".odt":  _extract_docx,
    ".html": _extract_html,
    ".htm":  _extract_html,
    ".csv":  _extract_csv,
    ".txt":  _extract_plaintext,
    ".md":   _extract_plaintext,
}


def extract_text(filepath: Path) -> dict:
    """Extract plain text from *filepath*.

    Returns a result dict as described in the module docstring.
    Never raises — all errors are captured in result["error"].
    """
    filepath = Path(filepath)
    ext = filepath.suffix.lower()

    result = {
        "text":     "",
        "type":     ext.lstrip("."),
        "source":   filepath,
        "metadata": {},
        "error":    None,
        "is_scan":  False,
    }

    extractor = _EXTRACTORS.get(ext)
    if extractor is None:
        result["error"] = f"Unsupported file type: {ext!r}"
        return result

    try:
        if ext == ".pdf":
            # PDF extractor returns a 3-tuple; others return a 2-tuple.
            text, metadata, is_scan = extractor(filepath)
            result["is_scan"] = is_scan
        else:
            text, metadata = extractor(filepath)

        result["text"]     = text.strip()
        result["metadata"] = metadata
    except Exception as exc:
        result["error"] = f"{type(exc).__name__}: {exc}"

    return result


# ---------------------------------------------------------------------------
# Command-line interface
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python scripts/extract_text.py <filepath>")
        sys.exit(1)

    target = Path(sys.argv[1])
    if not target.exists():
        print(f"Error: file not found: {target}", file=sys.stderr)
        sys.exit(1)

    res = extract_text(target)

    if res["error"]:
        print(f"Extraction error: {res['error']}", file=sys.stderr)
        sys.exit(1)

    if res["is_scan"]:
        print(
            f"[SCAN] {target.name} appears to be a scanned PDF with no text layer. "
            "Re-route to ocr.py.",
            file=sys.stderr,
        )
        sys.exit(2)

    if res["metadata"]:
        print("--- metadata ---", file=sys.stderr)
        for k, v in res["metadata"].items():
            print(f"  {k}: {v}", file=sys.stderr)

    print(res["text"])
